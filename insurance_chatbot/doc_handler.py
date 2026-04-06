from PyPDF2 import PdfReader  # Import PdfReader to read PDF files
from docx import Document  # Import Document to read DOCX files
from fpdf import FPDF  # For creating PDFs (not used for text extraction)
from datetime import datetime
import markdown
import streamlit as st
import pdfkit
import io

def process_doc(file_content, file_type):
    try:
        # Extract text from the file
        if file_type == "application/pdf":
            pdf_reader = PdfReader(file_content)
            if not pdf_reader.pages:  # Check if the PDF has pages
                return [False, "The uploaded PDF file is empty or invalid."]
            
            # Extract text from each page using PyPDF2
            page_texts = []
            for page_num, page in enumerate(pdf_reader.pages, start=1):
                text = page.extract_text().strip()
                if text:
                    page_texts.append({"page_number": page_num, "text": text})
            if not page_texts:  # Check if any text was extracted
                return [False, "The uploaded PDF file contains no readable text."]
            return [True, page_texts]

        elif file_type == "text/plain":
            text = file_content.getvalue().decode("utf-8")
            if not text.strip():  # Check if the TXT file is empty
                return [False, "The uploaded TXT file is empty."]
            return [True, [{"page_number": 1, "text": text}]]  # Treat as a single page

        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(file_content)
            if not doc.paragraphs:  # Check if the DOCX file has paragraphs
                return [False, "The uploaded DOCX file is empty or invalid."]
            
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            if not text.strip():  # Check if the extracted text is empty
                return [False, "The uploaded DOCX file contains no readable text."]
            return [True, [{"page_number": 1, "text": text}]]  # Treat as a single page

        else:
            return [False, "Unsupported file type. Please upload a valid PDF, TXT, or DOCX file."]

    except Exception as e:
        return [False, f"An error occurred while processing the document: {e}"]
    
def generate_analysis_pdf(insurance_type, user_data, claims_data, premium_data):
    """Generate a styled PDF report with analysis results."""
    class PDF(FPDF):
        def header(self):
            self.set_font('Arial', 'B', 12)
            self.cell(0, 10, f'🔍 Insurance Analysis Report - {insurance_type}', 0, 1, 'C')
            self.ln(5)
            
        def footer(self):
            self.set_y(-15)
            self.set_font('Arial', 'I', 8)
            self.cell(0, 10, f'Generated on {datetime.now().strftime("%Y-%m-%d")}', 0, 0, 'C')
    
    pdf = PDF()
    pdf.add_page()
    
    # User Details
    pdf.set_font('Arial', 'B', 12)
    pdf.cell(0, 10, '📋 User Details', 0, 1)
    pdf.set_font('Arial', '', 10)
    for key, value in user_data.items():
        pdf.cell(0, 8, f'• {key}: {value}', 0, 1)
    pdf.ln(5)
    
    # Claims Analysis
    pdf.set_font('Arial', 'B', 12)
    pdf.cell(0, 10, '📊 Claims Analysis', 0, 1)
    pdf.set_font('Arial', '', 10)
    for metric, value in zip(claims_data['Metric'], claims_data['Value']):
        pdf.cell(0, 8, f'• {metric}: {value}', 0, 1)
    pdf.ln(5)
    
    # Premium Calculation
    pdf.set_font('Arial', 'B', 12)
    pdf.cell(0, 10, '💰 Premium Calculation', 0, 1)
    pdf.set_font('Arial', '', 10)
    for component, amount in zip(premium_data['Component'], premium_data['Amount']):
        pdf.cell(0, 8, f'• {component}: {amount}', 0, 1)
    
    # Convert to bytes buffer
    buffer = io.BytesIO()
    pdf_bytes = pdf.output(dest='S')  # Output PDF as a bytes object
    buffer.write(pdf_bytes)  # Write the bytes to the buffer
    buffer.seek(0)
    return buffer

def markdown_to_pdf(markdown_text, output_filename="output.pdf"):
    """
    Convert Markdown text to PDF and return a BytesIO buffer for downloading.
    
    Args:
        markdown_text (str): The Markdown content to convert.
        output_filename (str): Name of the output PDF file for download (default: 'output.pdf').
    
    Returns:
        BytesIO: PDF file as a stream for download, or None if an error occurs.
    """
    try:
        # Specify the path to wkhtmltopdf executable
        wkhtmltopdf_path = None
        # if platform.system() == "Windows":
        #     wkhtmltopdf_path = r"C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe"  # Adjust if installed elsewhere
        # else:  # Linux/Mac
        wkhtmltopdf_path = r"C:\practice\wkhtmltox-0.12.6-1.mxe-cross-win64\wkhtmltox\bin\wkhtmltopdf.exe"  # Adjust based on `which wkhtmltopdf` output

        # Configure pdfkit with the wkhtmltopdf path
        config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path)
        
        # Convert Markdown to HTML with table extension
        html = markdown.markdown(markdown_text, extensions=['tables'])
        
        # Wrap HTML in enhanced styling for better PDF rendering, especially tables
        html_content = f"""
        <html>
        <head>
            <style>
                body {{ 
                    font-family: Arial, sans-serif; 
                    margin: 20px; 
                    line-height: 1.4; 
                }}
                h1 {{ 
                    color: #333; 
                    font-size: 24px; 
                    border-bottom: 2px solid #333; 
                    padding-bottom: 5px; 
                }}
                h2, h3 {{ 
                    color: #444; 
                    font-size: 18px; 
                    margin-top: 20px; 
                }}
                table {{ 
                    border-collapse: collapse; 
                    width: 100%; 
                    margin: 10px 0; 
                    background-color: #fff; 
                    box-shadow: 0 1px 3px rgba(0,0,0,0.1); 
                }}
                th, td {{ 
                    border: 1px solid #999; 
                    padding: 10px; 
                    text-align: left; 
                    font-size: 14px; 
                }}
                th {{ 
                    background-color: #f2f2f2; 
                    font-weight: bold; 
                    color: #333; 
                }}
                td {{ 
                    vertical-align: top; 
                }}
                code {{ 
                    background: #f4f4f4; 
                    padding: 2px 4px; 
                    border-radius: 4px; 
                    font-family: 'Courier New', Courier, monospace; 
                }}
                ul, ol {{ 
                    margin: 10px 0; 
                    padding-left: 20px; 
                }}
                p {{ 
                    margin: 8px 0; 
                }}
            </style>
        </head>
        <body>{html}</body>
        </html>
        """
        
        # Convert HTML to PDF bytes
        pdf_bytes = pdfkit.from_string(html_content, False, configuration=config, options={
            "encoding": "UTF-8",
            "page-size": "A4",
            "margin-top": "20mm",
            "margin-bottom": "20mm",
            "margin-left": "15mm",
            "margin-right": "15mm"
        })
        
        # Store PDF bytes in BytesIO
        pdf_output = io.BytesIO()
        pdf_output.write(pdf_bytes)
        pdf_output.seek(0)
        
        return pdf_output
    
    except Exception as e:
        return (f"Error converting Markdown to PDF: {str(e)}")
        

# Streamlit app
def main():
    pass

if __name__ == "__main__":
    main()
